#!/usr/bin/env python3
"""Generate E2E test data files for DS-1 ~ DS-5 acceptance datasets.

Usage:
    python tests/e2e/generate_test_data.py

Dependencies (NOT added to backend/requirements.txt):
    openpyxl, python-docx, reportlab, fpdf2, Pillow, pdfplumber
"""

import os
from pathlib import Path

import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.units import mm
import io

BASE_DIR = Path(__file__).resolve().parent / "test-data"

# ---------------------------------------------------------------------------
# Helper: create xlsx with auto column width
# ---------------------------------------------------------------------------

def create_xlsx(filepath: str, sheet_name: str, headers: list[str], rows: list[list]):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet_name
    ws.append(headers)
    for row in rows:
        ws.append(row)
    # auto width
    for col_idx, _ in enumerate(headers, 1):
        max_len = max(
            len(str(ws.cell(row=r, column=col_idx).value or ""))
            for r in range(1, ws.max_row + 1)
        )
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 40)
    wb.save(filepath)
    print(f"  [xlsx] {filepath}")


# ---------------------------------------------------------------------------
# DS-1: Standard scenario - 3 xlsx files
# ---------------------------------------------------------------------------

def gen_ds1():
    d = BASE_DIR / "ds-1"
    d.mkdir(parents=True, exist_ok=True)

    # Supplier A
    create_xlsx(
        str(d / "联想科技-办公设备报价.xlsx"),
        "报价明细",
        ["序号", "品名", "规格型号", "单位", "数量", "单价", "金额", "税率", "交货期", "备注"],
        [
            [1, "ThinkPad E14 笔记本", "i5-1335U/16G/512G SSD", "台", 50, 4299, 214950, "13%", "7个工作日", "含3年上门保修"],
            [2, "联想 27 寸显示器", "L27e-40 4K IPS", "台", 50, 1899, 94950, "13%", "5个工作日", ""],
            [3, "联想无线键鼠套装", "EC200", "套", 50, 129, 6450, "13%", "3个工作日", ""],
            [4, "ThinkPad 扩展坞", "USB-C Gen2", "个", 30, 899, 26970, "13%", "7个工作日", ""],
            [5, "联想激光打印机", "LJ2655DN", "台", 5, 2199, 10995, "13%", "5个工作日", "含1年保修"],
        ],
    )

    # Supplier B
    create_xlsx(
        str(d / "华硕经销-办公设备报价单.xlsx"),
        "Sheet1",
        ["产品名称", "规格", "计量单位", "采购数量", "报价", "合计", "增值税率", "供货周期", "说明"],
        [
            ["ASUS VivoBook 14 笔记本", "i5-1335U/16G/512G", "台", 50, 3899, 194950, "13%", "10个工作日", "含2年保修"],
            ["ASUS ProArt 27寸显示器", "PA279CRV 4K IPS", "台", 50, 3299, 164950, "13%", "7个工作日", "专业级色准"],
            ["华硕无线键鼠套装", "CW100", "套", 50, 99, 4950, "13%", "3个工作日", ""],
            ["华硕扩展坞", "USB-C DC300", "个", 30, 799, 23970, "13%", "5个工作日", ""],
            ["惠普激光打印机", "LaserJet M208dw", "台", 5, 1899, 9495, "13%", "5个工作日", "含1年保修"],
        ],
    )

    # Supplier C
    create_xlsx(
        str(d / "戴尔直销-报价函.xlsx"),
        "报价表",
        ["商品名称", "配置", "单位", "数量", "含税单价", "总价", "税率", "交期", "备注"],
        [
            ["Dell Latitude 3440 笔记本电脑", "i5-1335U/16G/512G SSD", "台", 50, 4099, 204950, "13%", "5-7个工作日", "3年ProSupport"],
            ["戴尔 U2723QE 27寸显示器", "4K IPS USB-C", "台", 50, 2899, 144950, "13%", "3-5个工作日", ""],
            ["戴尔无线键鼠套装", "KM5221W", "套", 50, 159, 7950, "13%", "3个工作日", ""],
            ["戴尔扩展坞", "WD19DCS USB-C", "个", 30, 1299, 38970, "13%", "7-10个工作日", ""],
            ["兄弟激光打印机", "HL-L2350DW", "台", 5, 1599, 7995, "13%", "5个工作日", "含1年保修"],
        ],
    )
    print("DS-1 done.\n")


# ---------------------------------------------------------------------------
# DS-2: Mixed formats (xlsx + docx + pdf)
# ---------------------------------------------------------------------------

def gen_ds2():
    d = BASE_DIR / "ds-2"
    d.mkdir(parents=True, exist_ok=True)

    # Supplier A - xlsx
    create_xlsx(
        str(d / "优品办公-家具报价.xlsx"),
        "报价",
        ["物资名称", "参数", "计量", "需求量", "价格", "小计", "税点", "货期", "附注"],
        [
            ["办公桌 1.4 米", "橡木/钢架 120×60×75cm", "张", 50, 1280, 64000, "13%", "15天", "含安装"],
            ["人体工学椅", "网布/可调节扶手/头枕", "把", 50, 1580, 79000, "13%", "10天", "5年质保"],
            ["文件柜三层", "钢制/带锁/灰白色", "组", 30, 680, 20400, "13%", "7天", ""],
            ["会议桌 3.6 米", "橡木/16人位", "张", 2, 8900, 17800, "13%", "20天", "含配套线槽"],
            ["屏风隔断 1.2 米", "铝合金框/布艺", "块", 40, 420, 16800, "13%", "10天", ""],
        ],
    )

    # Supplier B - docx
    gen_ds2_docx(d)

    # Supplier C - digital PDF
    gen_ds2_pdf(d)

    print("DS-2 done.\n")


def gen_ds2_docx(d: Path):
    doc = Document()
    doc.add_paragraph("致贵公司：以下为我司办公家具报价，有效期 30 天。")

    headers = ["品名", "规格型号", "单位", "数量", "含税单价（元）", "含税合计（元）", "增值税", "交货周期", "备注"]
    rows = [
        ["1.4m 办公桌", "胡桃木/钢架 120×60×75", "张", "50", "1150", "57500", "13%", "12天", "含安装"],
        ["人体工学转椅", "网布/4D扶手/头枕/腰靠", "把", "50", "1880", "94000", "13%", "14天", "3年质保"],
        ["三层文件柜", "冷轧钢/密码锁/白色", "组", "30", "720", "21600", "13%", "5天", ""],
        ["大型会议桌", "胡桃木/可坐16人/含线盒", "张", "2", "9500", "19000", "13%", "25天", ""],
        ["办公屏风 1.2m", "铝框/灰色布面", "块", "40", "380", "15200", "13%", "8天", ""],
    ]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows, 1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val

    filepath = str(d / "佳美家具报价函.docx")
    doc.save(filepath)
    print(f"  [docx] {filepath}")


def _find_cjk_font() -> str | None:
    """Try to find a CJK font on macOS / Linux."""
    candidates = [
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/Supplemental/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None


def _register_reportlab_cjk_font():
    """Register a CJK font for reportlab if available."""
    font_path = _find_cjk_font()
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont("CJK", font_path, subfontIndex=0))
            return "CJK"
        except Exception:
            pass
    # fallback: use reportlab's built-in CJK support
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    return "STSong-Light"


def gen_ds2_pdf(d: Path):
    filepath = str(d / "宏远办公-报价单.pdf")
    font_name = _register_reportlab_cjk_font()

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(filepath, pagesize=A4)
    elements = []

    # Title
    title_style = styles["Title"].clone("title_cjk")
    title_style.fontName = font_name
    title_style.fontSize = 16
    elements.append(Paragraph("宏远办公用品有限公司", title_style))
    elements.append(Spacer(1, 12))

    subtitle_style = styles["Normal"].clone("sub_cjk")
    subtitle_style.fontName = font_name
    subtitle_style.fontSize = 12
    elements.append(Paragraph("报价单", subtitle_style))
    elements.append(Spacer(1, 12))

    headers = ["项目名称", "型号/规格", "单位", "数量", "不含税单价", "不含税金额", "税率", "交期", "说明"]
    rows_data = [
        ["办公桌1400mm", "松木/金属腿 120×60×75", "张", "50", "980", "49000", "13%", "10天", "含安装"],
        ["网布人体工学椅", "三级气杆/调节扶手", "把", "50", "1200", "60000", "13%", "12天", "2年质保"],
        ["钢制文件柜(三层)", "0.8mm冷轧钢/暗锁", "组", "30", "550", "16500", "13%", "7天", ""],
        ["16人会议桌", "松木/3600×1500mm", "张", "2", "7200", "14400", "13%", "18天", "不含线槽"],
        ["屏风隔断1200mm", "铝合金/布艺/灰色", "块", "40", "350", "14000", "13%", "10天", ""],
    ]

    table_data = [headers] + rows_data
    col_widths = [80, 100, 30, 30, 55, 55, 30, 40, 50]

    t = Table(table_data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("ALIGN", (2, 1), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements.append(t)
    doc.build(elements)
    print(f"  [pdf ] {filepath}")


# ---------------------------------------------------------------------------
# DS-3: Anomaly scenario - 4 xlsx files
# ---------------------------------------------------------------------------

def gen_ds3():
    d = BASE_DIR / "ds-3"
    d.mkdir(parents=True, exist_ok=True)

    # Supplier A
    create_xlsx(
        str(d / "科达试剂-报价.xlsx"),
        "报价明细",
        ["品名", "规格", "单位", "数量", "含税单价", "总价", "税率", "交期"],
        [
            ["无水乙醇", "AR 500ml/瓶", "瓶", 100, 28, 2800, "13%", "3天"],
            ["浓硫酸", "AR 500ml/瓶", "瓶", 50, 35, 1750, "13%", "3天"],
            ["电子天平", "0.001g/220g", "台", 2, 12800, 25600, "13%", "7天"],
            ["移液器", "100-1000μl", "支", 10, 2800, 28000, "13%", "5天"],
            ["离心管", "50ml 无菌", "袋(25支/袋)", 20, 45, 900, "13%", "3天"],
        ],
    )

    # Supplier B
    create_xlsx(
        str(d / "国药试剂-报价单.xlsx"),
        "报价明细",
        ["商品名称", "规格型号", "单位", "数量", "不含税单价", "不含税金额", "税率", "交期", "备注"],
        [
            ["分析纯无水乙醇", "500ml", "瓶", 100, 22, 2200, "13%", "2天", ""],
            ["分析纯浓硫酸", "500ml", "瓶", 50, 30, 1500, "13%", "2天", ""],
            ["精密电子天平", "0.001g/200g", "台", 2, 11500, 23000, "13%", "10天", ""],
            ["微量移液器", "100-1000μl 可调", "把", 10, 2500, 25000, "13%", "7天", ""],
            ["离心管 50ml", "无菌 25支装", "包", 20, 38, 760, "13%", "2天", ""],
        ],
    )

    # Supplier C - special: no tax column, missing fields, delivery in remarks
    create_xlsx(
        str(d / "西陇试剂-报价.xlsx"),
        "报价明细",
        ["产品名", "规格", "单位", "数量", "单价", "金额", "交期", "备注"],
        [
            ["乙醇(无水)", "AR 500ml", "瓶", 100, 25, 2500, "", "5天到货"],
            ["硫酸(浓)", "AR 500ml", "瓶", 50, 32, 1600, "", ""],
            ["电子分析天平", "万分之一/220g", "台", 2, 13500, 27000, "", "10-15天"],
            ["移液枪", "100-1000μl", "支", 10, None, None, "", "缺货"],
            ["离心管", "50ml/袋装", "箱(500支/箱)", 20, 380, 7600, "", "3天"],
        ],
    )

    # Supplier D - English, USD, only 2 rows
    create_xlsx(
        str(d / "Sigma-Aldrich-Quote.xlsx"),
        "Quote",
        ["Product Name", "Specification", "Unit", "Qty", "Unit Price (USD)", "Total (USD)", "Tax", "Delivery"],
        [
            ["Ethanol, absolute", "AR 500ml", "bottle", 100, 5.50, 550, "0%", "14 days"],
            ["Sulfuric acid, concentrated", "AR 500ml", "bottle", 50, 8.00, 400, "0%", "14 days"],
        ],
    )

    print("DS-3 done.\n")


# ---------------------------------------------------------------------------
# DS-4: Grouping challenge - 3 xlsx files with different column names
# ---------------------------------------------------------------------------

def gen_ds4():
    d = BASE_DIR / "ds-4"
    d.mkdir(parents=True, exist_ok=True)

    # Supplier A: 品名、规格型号、单位、数量、单价
    create_xlsx(
        str(d / "供应商A报价.xlsx"),
        "报价",
        ["品名", "规格型号", "单位", "数量", "单价"],
        [
            ["Cisco 千兆交换机", "C9200L-24T-4G", "台", 5, 15800],
            ["华为企业级无线AP", "AirEngine 5761-21", "台", 20, 3200],
            ["HP 服务器", "ProLiant DL380 Gen10/Xeon 4314/32G/1.2T SAS", "台", 2, 68000],
            ["Cisco 企业级路由器产品", "ISR 4331", "台", 2, 22000],
            ["海康威视网络摄像头设备", "DS-2CD2T47G2-L 4MP", "个", 30, 680],
            ["海康威视 NVR 录像机", "DS-7608NI-K2 8路", "台", 2, 1580],
        ],
    )

    # Supplier B: 商品名称、规格、单位、数量、报价
    create_xlsx(
        str(d / "供应商B报价.xlsx"),
        "报价",
        ["商品名称", "规格", "单位", "数量", "报价"],
        [
            ["思科24口千兆交换机", "Catalyst 9200L-24T-4G", "台", 5, 16200],
            ["HUAWEI 室内无线AP", "AirEngine 5761-21", "台", 20, 2980],
            ["惠普机架式服务器", "DL380 Gen10 Plus/Xeon-4314/32GB/1.2TB", "台", 2, 72000],
            ["思科集成多业务路由器", "ISR4331/K9", "台", 2, 21500],
            ["Hikvision 枪型网络摄像机", "2CD2T47G2-L 400万像素", "个", 30, 720],
            ["Hikvision 网络硬盘录像机", "7608NI-K2 八通道", "台", 2, 1650],
        ],
    )

    # Supplier C: 产品名称、型号/配置、单位、数量、含税单价
    create_xlsx(
        str(d / "供应商C报价.xlsx"),
        "报价",
        ["产品名称", "型号/配置", "单位", "数量", "含税单价"],
        [
            ["24口千兆管理型交换机", "C9200L-24T-4G-E", "台", 5, 16800],
            ["Huawei 无线接入点", "5761-21", "台", 20, 3100],
            ["惠普 ProLiant 服务器", "DL380G10/至强4314/32G/1.2T", "台", 2, 69500],
            ["企业路由器", "Cisco ISR 4331", "台", 2, 23000],
            ["海康4MP网络摄像头", "DS-2CD2T47G2-L", "个", 30, 650],
            ["海康8路NVR", "DS-7608NI-K2", "台", 2, 1520],
        ],
    )

    print("DS-4 done.\n")


# ---------------------------------------------------------------------------
# DS-5: OCR path - 1 xlsx + 1 scanned PDF
# ---------------------------------------------------------------------------

def gen_ds5():
    d = BASE_DIR / "ds-5"
    d.mkdir(parents=True, exist_ok=True)

    # Supplier A - standard xlsx
    create_xlsx(
        str(d / "标准供应商A.xlsx"),
        "报价",
        ["品名", "规格", "单位", "数量", "单价", "税率"],
        [
            ["A4 打印纸", "70g 500张/包", "包", 200, 18, "13%"],
            ["黑色签字笔", "0.5mm", "支", 500, 1.5, "13%"],
            ["文件袋", "A4 透明按扣", "个", 300, 0.8, "13%"],
        ],
    )

    # Supplier B - scanned PDF (image-only)
    gen_ds5_scanned_pdf(d)

    print("DS-5 done.\n")


def gen_ds5_scanned_pdf(d: Path):
    """Generate a scanned-style PDF (text rendered as image, not extractable by pdfplumber)."""
    filepath = str(d / "扫描版报价单.pdf")

    # Render table as image using Pillow
    width, height = 800, 400
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    # Try to find a suitable font
    font = None
    font_candidates = [
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]
    for fp in font_candidates:
        if os.path.exists(fp):
            try:
                font = ImageFont.truetype(fp, 16)
                break
            except Exception:
                continue
    if font is None:
        font = ImageFont.load_default()

    small_font = font

    # Title
    draw.text((250, 20), "供应商B 报价单", fill="black", font=font)

    # Table data
    headers = ["品名", "规格", "单位", "数量", "单价", "税率"]
    rows = [
        ["A4 打印纸", "70g 500张/包", "包", "200", "16", "13%"],
        ["黑色签字笔", "0.5mm", "支", "500", "1.8", "13%"],
        ["文件袋", "A4 透明按扣", "个", "300", "0.7", "13%"],
    ]

    col_widths = [120, 130, 50, 50, 60, 60]
    x_start, y_start = 50, 70
    row_height = 35

    # Draw header
    x = x_start
    for i, h in enumerate(headers):
        draw.rectangle([x, y_start, x + col_widths[i], y_start + row_height], outline="black", fill="#E0E0E0")
        draw.text((x + 5, y_start + 8), h, fill="black", font=small_font)
        x += col_widths[i]

    # Draw data rows
    for r_idx, row in enumerate(rows):
        y = y_start + (r_idx + 1) * row_height
        x = x_start
        for c_idx, val in enumerate(row):
            draw.rectangle([x, y, x + col_widths[c_idx], y + row_height], outline="black")
            draw.text((x + 5, y + 8), val, fill="black", font=small_font)
            x += col_widths[c_idx]

    # Add slight noise to simulate scan
    import random
    random.seed(42)
    pixels = img.load()
    for _ in range(500):
        rx = random.randint(0, width - 1)
        ry = random.randint(0, height - 1)
        pixels[rx, ry] = (
            random.randint(180, 220),
            random.randint(180, 220),
            random.randint(180, 220),
        )

    # Slight rotation to simulate scan misalignment
    img = img.rotate(0.5, fillcolor="white", expand=False)

    # Save image to buffer, embed in PDF as full-page image
    img_buffer = io.BytesIO()
    img.save(img_buffer, format="PNG")
    img_buffer.seek(0)

    # Use reportlab to create a PDF with just the image
    from reportlab.lib.pagesizes import A4 as rl_A4
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(filepath, pagesize=rl_A4)
    pw, ph = rl_A4

    # Save temp image
    tmp_img_path = str(d / "_tmp_scan.png")
    img.save(tmp_img_path)

    c.drawImage(tmp_img_path, 50, ph - height - 100, width=width * 0.7, height=height * 0.7)
    c.save()

    # Clean up temp
    os.remove(tmp_img_path)
    print(f"  [pdf ] {filepath} (scanned/image-only)")


# ---------------------------------------------------------------------------
# Verification
# ---------------------------------------------------------------------------

def verify():
    print("=" * 60)
    print("VERIFICATION")
    print("=" * 60)

    errors = []

    # Check all expected files exist
    expected_files = [
        "ds-1/联想科技-办公设备报价.xlsx",
        "ds-1/华硕经销-办公设备报价单.xlsx",
        "ds-1/戴尔直销-报价函.xlsx",
        "ds-2/优品办公-家具报价.xlsx",
        "ds-2/佳美家具报价函.docx",
        "ds-2/宏远办公-报价单.pdf",
        "ds-3/科达试剂-报价.xlsx",
        "ds-3/国药试剂-报价单.xlsx",
        "ds-3/西陇试剂-报价.xlsx",
        "ds-3/Sigma-Aldrich-Quote.xlsx",
        "ds-4/供应商A报价.xlsx",
        "ds-4/供应商B报价.xlsx",
        "ds-4/供应商C报价.xlsx",
        "ds-5/标准供应商A.xlsx",
        "ds-5/扫描版报价单.pdf",
    ]

    for f in expected_files:
        fp = BASE_DIR / f
        if fp.exists():
            size = fp.stat().st_size
            print(f"  OK  {f} ({size} bytes)")
        else:
            errors.append(f"MISSING: {f}")
            print(f"  FAIL {f} - NOT FOUND")

    # Verify xlsx readback
    print("\n--- XLSX readback verification ---")
    xlsx_files = [f for f in expected_files if f.endswith(".xlsx")]
    for f in xlsx_files:
        fp = BASE_DIR / f
        if not fp.exists():
            continue
        try:
            wb = openpyxl.load_workbook(str(fp))
            ws = wb.active
            row_count = ws.max_row - 1  # minus header
            print(f"  OK  {f}: sheet='{ws.title}', rows={row_count}")
            wb.close()
        except Exception as e:
            errors.append(f"XLSX READ FAIL: {f}: {e}")
            print(f"  FAIL {f}: {e}")

    # Verify DS-3 supplier C special data
    print("\n--- DS-3 Supplier C special fields ---")
    fp = BASE_DIR / "ds-3/西陇试剂-报价.xlsx"
    if fp.exists():
        wb = openpyxl.load_workbook(str(fp))
        ws = wb.active
        # Row 5 (index 5 = row 4 of data = 移液枪) should have None for price/amount
        price_val = ws.cell(row=5, column=5).value  # 单价
        amount_val = ws.cell(row=5, column=6).value  # 金额
        if price_val is None and amount_val is None:
            print(f"  OK  移液枪 row: price={price_val}, amount={amount_val} (correctly None)")
        else:
            errors.append(f"DS-3 C 移液枪: expected None/None, got {price_val}/{amount_val}")
            print(f"  FAIL 移液枪 row: price={price_val}, amount={amount_val}")
        # Check no tax column (headers should not contain 税率)
        headers_row = [ws.cell(row=1, column=c).value for c in range(1, ws.max_column + 1)]
        has_tax = any("税" in str(h) for h in headers_row if h)
        if not has_tax:
            print(f"  OK  No tax column in headers: {headers_row}")
        else:
            errors.append("DS-3 C: unexpected tax column found")
            print(f"  FAIL Tax column found in headers: {headers_row}")
        wb.close()

    # Verify DS-2 PDF with pdfplumber
    print("\n--- DS-2 PDF pdfplumber extraction ---")
    import pdfplumber
    pdf_path = BASE_DIR / "ds-2/宏远办公-报价单.pdf"
    if pdf_path.exists():
        try:
            with pdfplumber.open(str(pdf_path)) as pdf:
                page = pdf.pages[0]
                tables = page.extract_tables()
                if tables and len(tables) > 0:
                    print(f"  OK  Extracted {len(tables)} table(s), first table has {len(tables[0])} rows")
                    for row in tables[0][:2]:
                        print(f"       {row}")
                else:
                    text = page.extract_text()
                    if text and len(text) > 20:
                        print(f"  OK  No table extracted but text found ({len(text)} chars)")
                    else:
                        errors.append("DS-2 PDF: no tables or text extracted")
                        print(f"  WARN No tables extracted from DS-2 PDF")
        except Exception as e:
            errors.append(f"DS-2 PDF: {e}")
            print(f"  FAIL {e}")

    # Verify DS-5 scanned PDF is NOT extractable
    print("\n--- DS-5 scanned PDF (should NOT have extractable text) ---")
    scan_path = BASE_DIR / "ds-5/扫描版报价单.pdf"
    if scan_path.exists():
        try:
            with pdfplumber.open(str(scan_path)) as pdf:
                page = pdf.pages[0]
                text = page.extract_text()
                tables = page.extract_tables()
                if not text and not tables:
                    print(f"  OK  No extractable text or tables (image-only PDF)")
                elif text and len(text.strip()) < 5:
                    print(f"  OK  Minimal text extracted: '{text.strip()[:50]}' (effectively image-only)")
                else:
                    print(f"  WARN Some text extracted ({len(text or '')} chars) - may not be fully image-only")
                    if text:
                        print(f"       Text sample: '{text[:100]}'")
        except Exception as e:
            errors.append(f"DS-5 PDF: {e}")
            print(f"  FAIL {e}")

    # Summary
    print("\n" + "=" * 60)
    if errors:
        print(f"VERIFICATION COMPLETED WITH {len(errors)} ERROR(S):")
        for e in errors:
            print(f"  - {e}")
    else:
        print("ALL VERIFICATIONS PASSED")
    print("=" * 60)

    return len(errors) == 0


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print(f"Generating test data in: {BASE_DIR}\n")
    gen_ds1()
    gen_ds2()
    gen_ds3()
    gen_ds4()
    gen_ds5()
    print()
    ok = verify()
    exit(0 if ok else 1)
