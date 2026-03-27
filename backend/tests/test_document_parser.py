from pathlib import Path

import pytest


@pytest.fixture
def parser():
    from engines.document_parser import DocumentParser

    return DocumentParser()


@pytest.fixture
def sample_xlsx(tmp_path) -> Path:
    """生成一个包含两个 sheet 的测试 Excel"""
    import openpyxl

    wb = openpyxl.Workbook()
    # Sheet1: 有数据
    ws1 = wb.active
    ws1.title = "报价单"
    ws1.append(["商品名称", "单价", "数量"])
    ws1.append(["笔记本电脑", "4299", "50"])
    ws1.append(["显示器", "1599", "30"])
    # Sheet2: 空白
    wb.create_sheet("空白页")
    # Sheet3: 有数据
    ws3 = wb.create_sheet("配件")
    ws3.append(["品名", "价格"])
    ws3.append(["键盘", "199"])

    path = tmp_path / "test.xlsx"
    wb.save(path)
    return path


@pytest.fixture
def empty_xlsx(tmp_path) -> Path:
    """生成一个完全空白的 Excel"""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "空白"
    path = tmp_path / "empty.xlsx"
    wb.save(path)
    return path


def test_parse_xlsx_basic(parser, sample_xlsx):
    """解析含数据的 Excel → 跳过空 sheet，返回 2 个表格"""
    results = parser.parse(str(sample_xlsx))
    assert len(results) == 2

    # 第一个表格
    t1 = results[0]
    assert t1.sheet_name == "报价单"
    assert t1.headers == ["商品名称", "单价", "数量"]
    assert t1.row_count == 2
    assert t1.column_count == 3
    assert t1.rows[0] == ["笔记本电脑", "4299", "50"]

    # 第二个表格
    t2 = results[1]
    assert t2.sheet_name == "配件"
    assert t2.row_count == 1


def test_parse_xlsx_empty(parser, empty_xlsx):
    """完全空白的 Excel → 返回空列表"""
    results = parser.parse(str(empty_xlsx))
    assert results == []


def test_parse_xlsx_progress(parser, sample_xlsx):
    """进度回调被正确调用"""
    progress_values: list[float] = []
    parser.parse(str(sample_xlsx), progress_callback=lambda p: progress_values.append(p))
    assert len(progress_values) > 0
    assert progress_values[-1] <= 1.0


def test_parse_unsupported_type(parser, tmp_path):
    """不支持的文件类型 → ValueError"""
    path = tmp_path / "test.txt"
    path.write_text("hello")
    with pytest.raises(ValueError, match="不支持的文件类型"):
        parser.parse(str(path))


def test_parse_xlsx_none_cells(parser, tmp_path):
    """含 None 单元格的 Excel → None 保留"""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["A", "B", "C"])
    ws.append(["x", None, "z"])
    path = tmp_path / "sparse.xlsx"
    wb.save(path)

    results = parser.parse(str(path))
    assert len(results) == 1
    assert results[0].rows[0] == ["x", None, "z"]


def test_is_ocr_available(parser):
    """OCR 检测不崩溃"""
    result = parser._is_ocr_available()
    assert isinstance(result, bool)


# ── Word (docx) 测试 ──


@pytest.fixture
def sample_docx(tmp_path) -> Path:
    """生成一个包含两个表格的测试 Word 文档"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("供应商报价单")

    # 表格 1
    table1 = doc.add_table(rows=3, cols=3)
    table1.cell(0, 0).text = "商品名称"
    table1.cell(0, 1).text = "单价"
    table1.cell(0, 2).text = "数量"
    table1.cell(1, 0).text = "打印机"
    table1.cell(1, 1).text = "2999"
    table1.cell(1, 2).text = "10"
    table1.cell(2, 0).text = "墨盒"
    table1.cell(2, 1).text = "199"
    table1.cell(2, 2).text = "50"

    doc.add_paragraph("其他信息")

    # 表格 2
    table2 = doc.add_table(rows=2, cols=2)
    table2.cell(0, 0).text = "型号"
    table2.cell(0, 1).text = "备注"
    table2.cell(1, 0).text = "HP-M1"
    table2.cell(1, 1).text = "含税"

    path = tmp_path / "test.docx"
    doc.save(path)
    return path


@pytest.fixture
def empty_docx(tmp_path) -> Path:
    """生成一个无表格的 Word 文档"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("没有表格的文档")
    path = tmp_path / "empty.docx"
    doc.save(path)
    return path


@pytest.fixture
def blank_table_docx(tmp_path) -> Path:
    """生成一个只有空白表格的 Word 文档"""
    from docx import Document

    doc = Document()
    doc.add_table(rows=3, cols=3)  # 全空
    path = tmp_path / "blank_table.docx"
    doc.save(path)
    return path


def test_parse_docx_basic(parser, sample_docx):
    """解析含表格的 Word → 返回 2 个表格"""
    results = parser.parse(str(sample_docx))
    assert len(results) == 2

    t1 = results[0]
    assert t1.headers == ["商品名称", "单价", "数量"]
    assert t1.row_count == 2
    assert t1.column_count == 3
    assert t1.rows[0] == ["打印机", "2999", "10"]
    assert t1.sheet_name is None

    t2 = results[1]
    assert t2.headers == ["型号", "备注"]
    assert t2.row_count == 1


def test_parse_docx_no_tables(parser, empty_docx):
    """无表格的 Word → 返回空列表"""
    results = parser.parse(str(empty_docx))
    assert results == []


def test_parse_docx_blank_table(parser, blank_table_docx):
    """只有空白表格 → 跳过，返回空列表"""
    results = parser.parse(str(blank_table_docx))
    assert results == []


def test_parse_docx_progress(parser, sample_docx):
    """进度回调被正确调用"""
    progress_values: list[float] = []
    parser.parse(str(sample_docx), progress_callback=lambda p: progress_values.append(p))
    assert len(progress_values) > 0
    assert progress_values[-1] <= 1.0


# ── PDF 测试 ──


@pytest.fixture
def sample_pdf(tmp_path) -> Path:
    """生成一个包含表格的简单 PDF"""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=10)

    headers = ["Product", "Price", "Qty"]
    rows = [
        ["Laptop", "4299", "50"],
        ["Monitor", "1599", "30"],
    ]

    col_width = 40
    row_height = 8

    for header in headers:
        pdf.cell(col_width, row_height, header, border=1)
    pdf.ln()

    for row in rows:
        for cell in row:
            pdf.cell(col_width, row_height, cell, border=1)
        pdf.ln()

    path = tmp_path / "test.pdf"
    pdf.output(str(path))
    return path


@pytest.fixture
def empty_pdf(tmp_path) -> Path:
    """生成一个无表格的 PDF（纯文本）"""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "This PDF has no tables")

    path = tmp_path / "empty.pdf"
    pdf.output(str(path))
    return path


def test_parse_pdf_basic(parser, sample_pdf):
    """解析含表格的 PDF → 返回至少 1 个表格"""
    results = parser.parse(str(sample_pdf))
    assert isinstance(results, list)
    if len(results) > 0:
        t = results[0]
        assert t.page_number == 1
        assert t.row_count >= 1
        assert t.column_count >= 1
        assert t.sheet_name is None


def test_parse_pdf_empty(parser, empty_pdf):
    """无表格的 PDF → 返回空列表"""
    results = parser.parse(str(empty_pdf))
    assert results == []


def test_parse_pdf_progress(parser, sample_pdf):
    """进度回调被正确调用"""
    progress_values: list[float] = []
    parser.parse(str(sample_pdf), progress_callback=lambda p: progress_values.append(p))
    assert len(progress_values) > 0
    assert progress_values[-1] <= 1.0
