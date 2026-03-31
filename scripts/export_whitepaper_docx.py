"""
白皮书 Markdown → DOCX 导出脚本

用法：
    python scripts/export_whitepaper_docx.py

输入：docs/whitepaper/software-whitepaper.md
输出：dist/whitepaper/software-whitepaper.docx

依赖：python-docx（pip install python-docx）

支持的 Markdown 构型（仅限白皮书实际使用的子集）：
    # H1         → 文档标题段落（Normal + 加粗居中，置于文件最前）
    ## H2        → Heading 2
    ### H3       → Heading 3
    **text**     → 加粗行内格式
    *text*       → 斜体行内格式
    ![alt](path) → 嵌入本地图片（路径相对于 .md 文件所在目录）
    | … |        → 普通表格（首行为表头）
    - item       → 无序列表
    ---          → 分隔线（输出为空段落，视觉上与下文保持间距）
    普通段落     → Normal 段落，支持行内加粗/斜体
"""

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

# ── 路径配置 ──────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
MD_PATH = PROJECT_ROOT / "docs" / "whitepaper" / "software-whitepaper.md"
OUT_PATH = PROJECT_ROOT / "dist" / "whitepaper" / "software-whitepaper.docx"
IMG_BASE = MD_PATH.parent  # 图片路径相对于 .md 文件所在目录


# ── 行内格式解析 ──────────────────────────────────────────────────────────────

def add_inline_text(para, text: str):
    """将含有 **bold** / *italic* 的文本逐段添加到段落。"""
    # 匹配 **bold**、*italic* 或普通文本
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))")
    for m in pattern.finditer(text):
        if m.group(2):  # **bold**
            run = para.add_run(m.group(2))
            run.bold = True
        elif m.group(3):  # *italic*
            run = para.add_run(m.group(3))
            run.italic = True
        elif m.group(4):  # plain
            para.add_run(m.group(4))


# ── 表格解析 ──────────────────────────────────────────────────────────────────

def parse_table_block(lines: list[str]) -> list[list[str]]:
    """将原始 Markdown 表格行（含分隔行）解析为 rows × cells 二维列表。"""
    rows = []
    for line in lines:
        if re.match(r"^\s*\|[-| :]+\|\s*$", line):
            continue  # 跳过分隔行
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            if ci >= col_count:
                break
            cell = table.cell(ri, ci)
            para = cell.paragraphs[0]
            para.clear()
            add_inline_text(para, cell_text)
            if ri == 0:  # 表头加粗
                for run in para.runs:
                    run.bold = True


# ── 图片插入 ──────────────────────────────────────────────────────────────────

def add_image(doc: Document, rel_path: str):
    """解析图片相对路径并嵌入文档，宽度限制为 6 英寸。"""
    img_path = (IMG_BASE / rel_path).resolve()
    if not img_path.exists():
        doc.add_paragraph(f"[图片缺失: {rel_path}]")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(img_path), width=Inches(6))


# ── 主解析循环 ─────────────────────────────────────────────────────────────────

def build_document(md_text: str) -> Document:
    doc = Document()

    # 默认正文字体
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    lines = md_text.splitlines()
    i = 0

    # 找到第一个 H1 作为文档标题
    title_written = False

    while i < len(lines):
        line = lines[i]

        # ── 跳过目录内容块（"## 目录" 开始到下一个 "---"） ──────────────────
        if line.strip() == "## 目录":
            while i < len(lines) and lines[i].strip() != "---":
                i += 1
            i += 1  # 跳过 ---
            continue

        # ── H1 → 文档标题 ────────────────────────────────────────────────────
        if line.startswith("# ") and not title_written:
            title_text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(18)
            title_written = True
            i += 1
            continue

        # ── H3 ───────────────────────────────────────────────────────────────
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # ── H2 ───────────────────────────────────────────────────────────────
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # ── 分隔线 ───────────────────────────────────────────────────────────
        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        # ── 图片 ─────────────────────────────────────────────────────────────
        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if img_match:
            add_image(doc, img_match.group(2))
            i += 1
            continue

        # ── 表格块（连续的 | 开头行） ─────────────────────────────────────────
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table_block(table_lines)
            add_table(doc, rows)
            continue

        # ── 无序列表项 ───────────────────────────────────────────────────────
        if re.match(r"^[-*] ", line):
            item_text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_text(p, item_text)
            i += 1
            continue

        # ── 空行 ─────────────────────────────────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── 普通段落（含行内加粗/斜体） ──────────────────────────────────────
        p = doc.add_paragraph()
        add_inline_text(p, line.strip())
        i += 1

    return doc


# ── 入口 ──────────────────────────────────────────────────────────────────────

def main():
    if not MD_PATH.exists():
        print(f"错误：找不到输入文件 {MD_PATH}", file=sys.stderr)
        sys.exit(1)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    md_text = MD_PATH.read_text(encoding="utf-8")
    doc = build_document(md_text)
    doc.save(str(OUT_PATH))

    size_kb = OUT_PATH.stat().st_size // 1024
    print(f"✓ 导出完成：{OUT_PATH}  ({size_kb} KB)")


if __name__ == "__main__":
    os.chdir(PROJECT_ROOT)
    main()
